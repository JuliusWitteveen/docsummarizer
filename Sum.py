"""
Document Summarizer

This script provides functionalities to summarize documents in various formats including PDF, DOCX, RTF, and TXT.
It uses NLP techniques for text processing and clustering to generate concise summaries.

Dependencies: fitz (PyMuPDF), docx, striprtf, langchain, sklearn, numpy, etc.
Usage: Run this script with Python 3.8 or above. Ensure all dependencies are installed.

Author: Julius Witteveen
Last Updated: 15-12-2023
Version: 2.0
"""

import re
import os
import logging
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tkinter.font as tkFont
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from sklearn.cluster import KMeans
import numpy as np
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
from langchain.chat_models import ChatOpenAI
from langdetect import detect
from kneed import KneeLocator
import threading
from docx import Document  # for DOCX file handling
import fitz  # for PDF file handling
from striprtf.striprtf import rtf_to_text  # for RTF file handling
from translate import Translator  # for prompt translation

# Global variables
selected_file_path = None
progress = None
custom_prompt_area = None

# Asynchronous Initialization Function
def init_async_components():
    # Load any heavy components or perform time-consuming tasks here
    # Example: Initializing langchain components, pre-loading models, etc.
    pass

# Background Summarization Function
def start_summarization_thread(root):
    summarization_thread = threading.Thread(target=start_summarization, args=(root,))
    summarization_thread.start()

def start_summarization(root):
    global selected_file_path, custom_prompt_area
    api_key = get_api_key()
    if api_key and selected_file_path:
        custom_prompt_text = get_summary_prompt(selected_file_path, api_key)
        
        # Generate summary and handle GUI updates in a thread-safe manner
        summary, language, filename_without_ext = generate_summary(
            selected_file_path, api_key, custom_prompt_text, 
            lambda value: update_progress_bar(value, root)
        )
        
        if summary and language:
            # Save the summary file and update the GUI
            root.after(0, lambda: save_summary_file(summary, language, filename_without_ext))
        # Finalize process and update the progress bar
        root.after(0, lambda: finalize_process(root))

def update_progress_bar(value, root):
    def set_progress(value):
        progress['value'] = value
    root.after(0, lambda: set_progress(value))

def finalize_process(root):
    update_progress_bar(100, root)

# Initialize logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Set default prompt in English
default_prompt_en = """
Summarize the text concisely and directly without prefatory phrases. Focus on presenting its key points and main ideas, ensuring that essential details are accurately conveyed in a straightforward manner.
"""

# -------------------------------------------------------------------
# Helper Functions
# -------------------------------------------------------------------

def is_valid_file_path(path):
    # Regular expression to validate a basic file path
    pattern = r'^[a-zA-Z0-9_\-\\/:. ]+$'
    return re.match(pattern, path) and os.path.isfile(path)

def get_api_key(file_path=r'C:\Users\HP\Documents\Python\api_key.txt'):
    logging.info("Retrieving API key.")

    # Validate the file path
    if not is_valid_file_path(file_path):
        logging.error(f"Invalid file path: {file_path}")
        return None

    try:
        with open(file_path, 'r') as file:
            return file.read().strip()
    except FileNotFoundError:
        logging.error(f"API key file not found at {file_path}")
        return None
    except IOError as e:
        logging.error(f"Error reading the API key file: {e}")
        return None

def select_file():
    file_path = filedialog.askopenfilename(
        title="Select a Document",
        filetypes=[("PDF Files", "*.pdf"), ("Word Documents", "*.docx"), ("RTF Files", "*.rtf"), ("Text Files", "*.txt")])
    return file_path

# -------------------------------------------------------------------
# Language Detection and Prompt Selection
# -------------------------------------------------------------------
def translate_prompt(prompt_text, target_language):
    if target_language == "nl":
        translator = Translator(to_lang="nl")
        try:
            translated_text = translator.translate(prompt_text)
            return translated_text
        except Exception as e:
            logging.error(f"Error translating prompt: {e}")
            # Return the original English prompt in case of translation failure
            return prompt_text

    # Return the original English prompt for unsupported languages
    return prompt_text

def get_summary_prompt(file_path, api_key):
    text, _ = load_document(file_path)
    if not text:
        return None

    language = detect_language(text)
    if language == "nl":
        # Attempt to translate the prompt, but fall back to English if it fails
        translated_prompt = translate_prompt(default_prompt_en, language)
        return translated_prompt
    elif language == "en":
        return default_prompt_en

    # Default to English prompt if language is not supported or translation fails
    return default_prompt_en

# -------------------------------------------------------------------
# Main Summarization Logic
# -------------------------------------------------------------------

def load_document(file_path):
    """
    Load and extract text from a document with lazy loading of dependencies.
    """
    logging.info(f"Loading document from: {file_path}")

    if not is_valid_file_path(file_path):
        logging.error(f"Invalid file path: {file_path}")
        return None, None

    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if file_extension == ".pdf":
            import fitz  # Lazy loading of fitz (PyMuPDF)
            with fitz.open(file_path) as doc:
                text = "\n".join(page.get_text() for page in doc)

        elif file_extension == ".docx":
            from docx import Document  # Lazy loading of docx
            doc = Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        elif file_extension == ".rtf":
            from striprtf.striprtf import rtf_to_text  # Lazy loading of striprtf
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_text = file.read()
            text = rtf_to_text(rtf_text)

        elif file_extension == ".txt":
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")

    except Exception as e:
        logging.error(f"Error loading the document: {e}")
        return None, None

    filename_without_ext = os.path.splitext(os.path.basename(file_path))[0]
    return text, filename_without_ext

def detect_language(text):
    """
    Detect the language of the text.
    """
    logging.info("Detecting language of the text.")
    try:
        return detect(text)
    except langdetect.lang_detect_exception.LangDetectException as e:
        logging.error(f"Language detection failed: {e}")
        return "unknown"

def find_elbow_point(sse):
    k_values = range(1, len(sse) + 1)
    kneedle = KneeLocator(k_values, sse, curve='convex', direction='decreasing')
    return kneedle.elbow if kneedle.elbow is not None else 20  # Default to 20 clusters if no elbow point is found

def prepare_text_for_summarization(text, openai_api_key):
    docs, vectors = split_and_embed_text(text, openai_api_key)
    num_chunks = len(docs)
    num_clusters = determine_optimal_clusters(vectors, max_clusters=min(num_chunks, 100))
    
    if num_clusters is None:
        num_clusters = min(num_chunks, 20)  # Default to a smaller number of clusters if None is returned

    if num_chunks < num_clusters:
        logging.warning(f"Only {num_chunks} chunks available, reducing the number of clusters to {num_chunks}.")
        num_clusters = num_chunks
    else:
        logging.info(f"Using {num_clusters} clusters out of {num_chunks} available text chunks based on the elbow point.")

    return docs, vectors, num_clusters

def determine_optimal_clusters(vectors, max_clusters=100):
    sse = []
    for k in range(1, max_clusters + 1):
        kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
        kmeans.fit(vectors)
        sse.append(kmeans.inertia_)

    elbow_point = find_elbow_point(sse)
    logging.info(f"Elbow method determined the optimal number of clusters as: {elbow_point}")
    return elbow_point

def split_and_embed_text(text, openai_api_key):
    """
    Split the text into manageable chunks and embed these chunks.
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(separators=["\n\n", "\n", "\t"], chunk_size=10000, chunk_overlap=3000)
        docs = text_splitter.create_documents([text])
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        vectors = embeddings.embed_documents([x.page_content for x in docs])
        return docs, vectors
    except Exception as e:
        logging.error(f"Error during text splitting and embedding: {e}")
        raise

def perform_clustering(vectors, max_clusters=100):
    """
    Determine the optimal number of clusters and find central points.
    """
    sse = []
    for k in range(1, max_clusters + 1):
        kmeans = KMeans(n_clusters=k, random_state=42)
        kmeans.fit(vectors)
        sse.append(kmeans.inertia_)

    elbow_point = find_elbow_point(sse)
    return elbow_point

from concurrent.futures import ThreadPoolExecutor, as_completed

def generate_chunk_summaries(docs, selected_indices, openai_api_key, custom_prompt, max_workers=10):
    """
    Generate summaries for each selected document chunk in parallel.
    """
    llm3_turbo = ChatOpenAI(temperature=0, openai_api_key=openai_api_key, max_tokens=4096, model='gpt-3.5-turbo-16k')
    map_prompt_template = PromptTemplate(template=f"```{{text}}```\n{custom_prompt}", input_variables=["text"])
    summary_list = []

    # Function to process each document chunk
    def process_chunk(doc):
        return load_summarize_chain(llm=llm3_turbo, chain_type="stuff", prompt=map_prompt_template).run([doc])

    # Parallel processing using ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_doc = {executor.submit(process_chunk, docs[i]): i for i in selected_indices}
        for future in as_completed(future_to_doc):
            index = future_to_doc[future]
            try:
                chunk_summary = future.result()
                if index < len(selected_indices) - 1:
                    chunk_summary += "\n"
                summary_list.append(chunk_summary)
            except Exception as e:
                logging.error(f"Error summarizing document chunk at index {index}: {e}")

    return "\n".join(summary_list)

def cluster_embeddings(vectors, num_clusters, update_progress):
    update_progress(30)
    logging.info(f"Clustering embeddings into {num_clusters} clusters.")
    try:
        kmeans = KMeans(n_clusters=num_clusters, random_state=42, n_init=10).fit(vectors)
        closest_indices = [np.argmin(np.linalg.norm(vectors - center, axis=1)) for center in kmeans.cluster_centers_]
        return sorted(closest_indices)
    except ValueError as e:
        logging.error(f"Error during clustering embeddings: {e}")
        raise

def compile_final_summary(summaries):
    """
    Compile individual summaries into one final summary.
    """
    final_summary = ''.join(summary for summary in summaries)
    return final_summary

# Sub-Function: Load Document and Detect Language
def load_document_and_detect_language(file_path):
    text, filename_without_ext = load_document(file_path)
    if not text:
        return None, None, "Failed to load document."

    language = detect_language(text)
    if language not in ['nl', 'en']:
        return None, None, f"Unsupported language detected: {language}"

    return text, filename_without_ext, None

# Sub-Function: Prepare Text for Summarization
def prepare_text_for_summarization(text, openai_api_key):
    docs, vectors = split_and_embed_text(text, openai_api_key)
    num_chunks = len(docs)
    num_clusters = determine_optimal_clusters(vectors, max_clusters=min(num_chunks, 100))
    
    if num_chunks < num_clusters:
        logging.warning(f"Only {num_chunks} chunks available, reducing the number of clusters to {num_chunks}.")
        num_clusters = num_chunks
    else:
        logging.info(f"Using {num_clusters} clusters out of {num_chunks} available text chunks based on the elbow point.")

    return docs, vectors, num_clusters

# Sub-Function: Generate and Compile Summary
def generate_and_compile_summary(docs, vectors, num_clusters, openai_api_key, custom_prompt, update_progress):
    selected_indices = cluster_embeddings(vectors, num_clusters, update_progress)
    summary = generate_chunk_summaries(docs, selected_indices, openai_api_key, custom_prompt)
    final_summary = compile_final_summary([summary])

    return final_summary

# Main Function: generate_summary
def generate_summary(file_path, openai_api_key, custom_prompt, update_progress):
    text, filename_without_ext, error_message = load_document_and_detect_language(file_path)
    if error_message:
        return error_message, None, filename_without_ext

    docs, vectors, num_clusters = prepare_text_for_summarization(text, openai_api_key)

    final_summary = generate_and_compile_summary(docs, vectors, num_clusters, openai_api_key, custom_prompt, update_progress)

    return final_summary, detect_language(text), filename_without_ext

# -------------------------------------------------------------------
# File Saving Logic
# -------------------------------------------------------------------

def save_summary(summary, file_path):
    """
    Save the summary in the format specified by the file_path extension.
    """
    try:
        if file_path.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(summary)

        elif file_path.endswith('.docx'):
            from docx import Document  # Lazy loading of docx
            doc = Document()
            doc.add_paragraph(summary)
            doc.save(file_path)

        elif file_path.endswith('.pdf'):
            from reportlab.platypus import SimpleDocTemplate, Paragraph  # Lazy loading of reportlab
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.pagesizes import letter

            pdf_file_path = file_path
            pdf = SimpleDocTemplate(pdf_file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            summary_paragraph = Paragraph(summary, styles['Normal'])
            pdf.build([summary_paragraph])

        else:
            raise ValueError("Unsupported file format selected.")

    except Exception as e:
        logging.error(f"An error occurred while saving the file: {e}")
        raise RuntimeError(f"An error occurred while saving the file: {e}")

    logging.info(f"Summary saved successfully to {file_path}")

def save_summary_file(summary, language, filename_without_ext):
    default_summary_filename = f"{filename_without_ext}_sum"
    file_path = filedialog.asksaveasfilename(
        initialfile=default_summary_filename,
        filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf")],
        defaultextension=".txt"
    )
    if file_path:
        save_summary(summary, file_path)
        messagebox.showinfo("Success", f"Summary saved successfully to {file_path}")
    else:
        messagebox.showerror("Error", "No file path selected.")


# -------------------------------------------------------------------
# GUI Code Block
# -------------------------------------------------------------------
def main_gui():
    global selected_file_path, progress, custom_prompt_area

    logging.info("Initializing GUI for the Document Summarizer.")
    root = tk.Tk()
    root.title("Document Summarizer")
    root.state('zoomed')  # Full-screen window

    # Define colors, fonts, and styles
    primary_color = "#2E3F4F"
    secondary_color = "#4F5D75"
    text_color = "#E0FBFC"
    button_color = "#3F88C5"
    larger_font = ('Arial', 12)
    button_font = ('Arial', 10, 'bold')

    style = ttk.Style()
    style.theme_use('clam')

    # Configure the style of TButton
    style.configure('W.TButton', font=button_font, background=button_color, foreground=text_color)
    style.map('W.TButton', background=[('active', secondary_color)], foreground=[('active', text_color)])

    # Configure layout
    root.configure(bg=primary_color)
    root.grid_columnconfigure(0, weight=1)
    root.grid_rowconfigure(1, weight=1)

    # Progress bar
    progress = ttk.Progressbar(root, orient=tk.HORIZONTAL, length=300, mode='determinate')
    progress.grid(row=0, column=0, pady=10, padx=10, sticky='ew')

    # Customizable prompt box
    prompt_label = tk.Label(root, text="Customize the summarization prompt:", fg=text_color, bg=primary_color, font=larger_font)
    prompt_label.grid(row=1, column=0, pady=(10, 0), sticky='nw')

    custom_prompt_area = tk.Text(root, height=15, width=80, wrap="word", bd=2, font=larger_font)
    custom_prompt_area.grid(row=2, column=0, pady=10, padx=10, sticky='nsew')

    # Function for file selection
    def file_select():
        global selected_file_path
        selected_file_path = select_file()
        if selected_file_path:
            api_key = get_api_key()
            if api_key:
                try:
                    text, filename_without_ext = load_document(selected_file_path)
                    language = detect_language(text)
                    custom_prompt = default_prompt_en  # Use the default English prompt
                    if language == "nl":
                        # Try to translate the prompt to Dutch, fallback to English if it fails
                        dutch_prompt = translate_prompt(default_prompt_en, "nl")
                        custom_prompt = dutch_prompt if dutch_prompt else default_prompt_en

                    custom_prompt_area.delete("1.0", tk.END)
                    custom_prompt_area.insert(tk.END, custom_prompt)
                    progress['value'] = 0
                    summarize_button['state'] = 'normal'
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load document: {e}")
                    summarize_button['state'] = 'disabled'
            else:
                messagebox.showinfo("API Key Error", "API key not found. Please check the file path.")
                summarize_button['state'] = 'disabled'
        else:
            summarize_button['state'] = 'disabled'

    # Loading label
    loading_label = tk.Label(root, text="Processing...", fg=text_color, bg=primary_color, font=larger_font)

    # Select file button
    select_button = ttk.Button(root, text="Select Document", command=file_select, style='W.TButton')
    select_button.grid(row=3, column=0, pady=20, padx=10, sticky='ew')

    # Start summarization button
    summarize_button = ttk.Button(root, text="Start Summarization", style='W.TButton')
    summarize_button['command'] = lambda: start_summarization_thread(root, summarize_button)
    summarize_button.grid(row=4, column=0, pady=20, padx=10, sticky='ew')

    # Helper function for updating the progress bar
    def set_progress(value):
        progress['value'] = value

    # Start summarization thread with updated button behavior
    def start_summarization_thread(root, summarize_button):
        summarize_button.config(state='disabled', text='Summarizing...')
        loading_label.grid(row=5, column=0, pady=10)  # Show loading label
        threading.Thread(target=lambda: start_summarization(root, summarize_button, loading_label)).start()

    def start_summarization(root, summarize_button, loading_label):
        global selected_file_path, custom_prompt_area
        api_key = get_api_key()
        if api_key and selected_file_path:
            custom_prompt_text = custom_prompt_area.get("1.0", "end-1c")

            # Update progress bar to indicate the start of summarization
            root.after(0, lambda: set_progress(10))

            # Generate the summary
            summary, language, filename_without_ext = generate_summary(
                selected_file_path, api_key, custom_prompt_text, 
                lambda value: update_progress_bar(value, root)
            )

            # Check if a summary was successfully generated
            if summary and language:
                # Save the summary file and update the GUI
                root.after(0, lambda: save_summary_file(summary, language, filename_without_ext))

            # Finalize process and update the progress bar to completion
            root.after(0, lambda: finalize_process(root))

        # Reset the GUI elements outside the if condition
        root.after(0, lambda: summarize_button.config(state='normal', text='Start Summarization'))
        root.after(0, loading_label.grid_forget)  # Hide loading label when done

    # Finalize the process and update the progress bar
    def finalize_process(root):
        set_progress(100)
        # Include any other clean-up or finalization steps here

    root.mainloop()  # This line starts the Tkinter event loop

# -------------------------------------------------------------------
# Script Execution Block
# -------------------------------------------------------------------

if __name__ == '__main__':
    logging.info("Starting the Document Summarizer application.")

    # Load environment variables only if necessary
    # For instance, you can load them before API key retrieval or summarization
    # load_dotenv()

    main_gui()